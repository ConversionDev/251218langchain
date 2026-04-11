"""
문서 추출 공통 모듈 (PDF / TXT / Word / HWP / Excel).

- 백엔드·프론트 모두 확장자 및 추출 로직은 여기서 관리. (프론트는 GET /api/document/supported-extensions 참고)
- 텍스트: extract_text_from_document() → PDF, TXT, Word .docx, HWP .hwp → str
- Excel: extract_excel_from_document() → .xlsx → list[dict] (첫 행 헤더, 시트 행을 dict로).
- PDF: domain/shared 전략(경로) 또는 bytes 시 pymupdf → pdfplumber.
- TXT: utf-8/cp949 디코딩.
- Word: python-docx (.docx만).
- HWP: pyhwp(hwp5txt) (.hwp v5).
- Excel: openpyxl (.xlsx만).
"""

import io
import logging
import tempfile
from contextlib import closing
from pathlib import Path
from typing import Any, Dict, List, Optional

logger = logging.getLogger(__name__)

# 지원 확장자 (단일 소스 — 백엔드 검증·API·프론트 accept 동기화용)
SUPPORTED_TEXT_EXTENSIONS = (".pdf", ".txt", ".docx", ".hwp")
SUPPORTED_EXCEL_EXTENSIONS = (".xlsx",)
SUPPORTED_EXTENSIONS = (*SUPPORTED_TEXT_EXTENSIONS, *SUPPORTED_EXCEL_EXTENSIONS)


def _extract_txt_from_bytes(data: bytes) -> str:
    """TXT: bytes 디코딩 (utf-8 우선, 실패 시 cp949)."""
    try:
        return data.decode("utf-8", errors="replace")
    except Exception:
        return data.decode("cp949", errors="replace")


def _extract_pdf_from_bytes(data: bytes) -> str:
    """PDF: bytes → pymupdf 우선, 실패 시 pdfplumber."""
    try:
        import fitz  # type: ignore[import-untyped]  # pymupdf

        doc = fitz.open(stream=data, filetype="pdf")
        try:
            parts = [page.get_text() for page in doc]
            return "\n".join(parts).strip() or "(PDF에서 텍스트를 추출하지 못했습니다)"
        finally:
            doc.close()
    except ImportError:
        pass
    except Exception as e:
        logger.warning("pymupdf PDF 추출 실패: %s", e)

    try:
        import pdfplumber

        with pdfplumber.open(io.BytesIO(data)) as doc:
            parts = [p.extract_text() or "" for p in doc.pages]
        return "\n".join(parts).strip() or "(PDF에서 텍스트를 추출하지 못했습니다)"
    except ImportError:
        raise RuntimeError("PDF 추출을 위해 pymupdf 또는 pdfplumber가 필요합니다.")
    except Exception as e2:
        raise RuntimeError(f"PDF 텍스트 추출 실패: {e2}")


def _extract_docx_from_bytes(data: bytes) -> str:
    """Word(.docx): python-docx로 본문 추출."""
    try:
        from docx import Document  # type: ignore[import-untyped]  # python-docx

        doc = Document(io.BytesIO(data))
        parts = [p.text for p in doc.paragraphs]
        # 표 내용도 포함
        for table in doc.tables:
            for row in table.rows:
                cells = [str(cell.text or "").strip() for cell in row.cells]
                if any(cells):
                    parts.append("\t".join(cells))
        return "\n".join(parts).strip() or "(Word에서 추출된 텍스트가 없습니다)"
    except ImportError:
        raise RuntimeError("Word(.docx) 추출을 위해 python-docx가 필요합니다. pip install python-docx")
    except Exception as e:
        raise RuntimeError(f"Word(.docx) 텍스트 추출 실패: {e}")


def _extract_pdf_from_path(path: Path) -> str:
    """PDF: 경로 → domain/shared StrategyFactory 사용 (disclosure 파이프라인 호환)."""
    from domain.shared.strategies.pdf_strategy import StrategyFactory  # type: ignore

    strategy = StrategyFactory.get_strategy(path)
    return strategy.extract(path)


def get_pdf_page_count(path: Path) -> int:
    """PDF 페이지 수 반환 (disclosure 검증 등). 경로만 지원."""
    try:
        import fitz  # type: ignore[import-untyped]  # pymupdf

        doc = fitz.open(path)
        try:
            return len(doc)
        finally:
            doc.close()
    except Exception as e:
        raise RuntimeError(f"PDF 페이지 수 확인 실패: {e}") from e


def _extract_docx_from_path(path: Path) -> str:
    """Word(.docx): 경로 → python-docx."""
    try:
        from docx import Document  # type: ignore[import-untyped]  # python-docx

        doc = Document(path)
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                cells = [str(cell.text or "").strip() for cell in row.cells]
                if any(cells):
                    parts.append("\t".join(cells))
        return "\n".join(parts).strip() or "(Word에서 추출된 텍스트가 없습니다)"
    except ImportError:
        raise RuntimeError("Word(.docx) 추출을 위해 python-docx가 필요합니다. pip install python-docx")
    except Exception as e:
        raise RuntimeError(f"Word(.docx) 텍스트 추출 실패: {e}")


def _try_extract_hwpx_zip(data: bytes) -> str:
    """확장자가 .hwp 이지만 내용이 ZIP인 HWPX(OOXML)인 경우 본문 XML에서 텍스트만 모은다.

    pyhwp(hwp5txt)는 **OLE 기반 HWP v5**만 처리한다. 한글 최근 버전의 HWPX는 여기서 처리."""
    import io
    import re
    import zipfile

    if not zipfile.is_zipfile(io.BytesIO(data)):
        return ""
    out: list[str] = []
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as zf:
            section_files = [
                n
                for n in zf.namelist()
                if re.search(r"(^|/)Contents/section\d+\.xml$", n.replace("\\", "/"), re.I)
            ]
            if not section_files:
                return ""
            for name in sorted(section_files):
                raw = zf.read(name)
                xml = raw.decode("utf-8", errors="replace")
                chunks = re.findall(r"<[^/\s>]+:t\b[^>]*>([^<]*)</[^>]+:t>", xml)
                if not chunks:
                    chunks = re.findall(r"<hp:t[^>]*>([^<]*)</hp:t>", xml)
                for c in chunks:
                    s = (
                        (c or "")
                        .replace("&lt;", "<")
                        .replace("&gt;", ">")
                        .replace("&amp;", "&")
                        .replace("&quot;", '"')
                        .strip()
                    )
                    if s:
                        out.append(s)
    except (zipfile.BadZipFile, OSError, KeyError) as e:
        logger.debug("HWPX(zip) 추출 생략: %s", e)
        return ""
    return "\n".join(out)


def _decode_hwp_pyhwp_output(raw: bytes) -> str:
    """pyhwp plaintext 변환 결과(bytes) → str. UTF-8/cp949 혼재에 대비."""
    if not raw:
        return ""

    def _hangul_chars(s: str) -> int:
        return sum(1 for c in s if "\uac00" <= c <= "\ud7a3")

    candidates: list[str] = []
    try:
        candidates.append(raw.decode("utf-8"))
    except UnicodeDecodeError:
        candidates.append(raw.decode("utf-8", errors="replace"))
    candidates.append(raw.decode("cp949", errors="replace"))
    candidates.append(raw.decode("euc-kr", errors="replace"))
    best = max(candidates, key=_hangul_chars)
    return best.strip()


def _extract_hwp_ole_pyhwp_path(path: Path) -> str:
    """OLE HWP v5: pyhwp events() API로 표 셀 포함 전체 텍스트 추출.

    기존 TextTransform은 표(Table) 내부 단락을 누락하는 한계가 있어,
    section.events()에서 'Text' 이벤트만 수집하는 방식으로 교체.
    표 셀·일반 단락 모두 동일한 이벤트 스트림으로 노출되므로 누락 없음.
    """
    try:
        from hwp5.errors import InvalidHwp5FileError  # type: ignore
        from hwp5.treeop import STARTEVENT  # type: ignore
        from hwp5.xmlmodel import Hwp5File  # type: ignore
    except ImportError as e:
        raise RuntimeError(
            "HWP(.hwp) 추출을 위해 pyhwp가 필요합니다. pip install pyhwp"
        ) from e

    import re

    parts: list[str] = []
    try:
        with closing(Hwp5File(str(path.resolve()))) as hwp5file:
            bt = hwp5file.bodytext
            for sec_idx in bt.section_indexes():
                sec = bt.section(sec_idx)
                for event, model in sec.events():
                    if event is not STARTEVENT:
                        continue
                    if not isinstance(model, tuple) or len(model) < 2:
                        continue
                    cls, data = model[0], model[1]
                    if getattr(cls, "__name__", "") == "Text" and isinstance(data, dict):
                        t = str(data.get("text") or "").strip()
                        if t:
                            parts.append(t)
    except InvalidHwp5FileError as e:
        raise RuntimeError(
            f"HWP 파일을 pyhwp로 열 수 없습니다. 호환 형식인지 확인하세요: {e}"
        ) from e
    except Exception as e:
        raise RuntimeError(f"HWP 본문 파싱 실패: {e}") from e

    text = "\n".join(parts)
    # 연속 공백/줄바꿈 정리
    text = re.sub(r"[ \t]{2,}", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = text.strip()

    if len(text) < 10:
        raise RuntimeError(
            "HWP에서 추출된 본문이 너무 짧습니다. "
            "스캔 이미지 위주 문서이거나 텍스트 레이어가 없을 수 있습니다."
        )
    return text


def _extract_hwp_from_bytes(data: bytes) -> str:
    """HWP(.hwp): HWPX(zip) 우선 → OLE v5는 pyhwp(hwp5txt).

    Windows: NamedTemporaryFile을 **닫은 뒤** pyhwp가 파일을 열도록 분리.
    (파일이 열려 있는 상태에서 pyhwp가 같은 경로를 열면 WinError 32 발생)
    """
    stripped = _try_extract_hwpx_zip(data).strip()
    if len(stripped) >= 10:
        return stripped

    tmp_path: Optional[Path] = None
    try:
        with tempfile.NamedTemporaryFile(suffix=".hwp", delete=False) as f:
            f.write(data)
            tmp_path = Path(f.name)
        # 파일 핸들이 완전히 닫힌 뒤 pyhwp 호출
        return _extract_hwp_ole_pyhwp_path(tmp_path)
    finally:
        if tmp_path is not None:
            tmp_path.unlink(missing_ok=True)


def _extract_hwp_from_path(path: Path) -> str:
    """HWP(.hwp): HWPX(ZIP) 우선 → OLE v5는 pyhwp in-process."""
    data = path.read_bytes()
    stripped = _try_extract_hwpx_zip(data).strip()
    if len(stripped) >= 10:
        return stripped
    return _extract_hwp_ole_pyhwp_path(path)


def extract_text_from_document(
    *,
    data: Optional[bytes] = None,
    path: Optional[Path] = None,
    filename: Optional[str] = None,
) -> str:
    """문서에서 텍스트 추출 (PDF / TXT / Word .docx / HWP .hwp).

    호출 시 (data + filename) 또는 path 중 하나만 넘기면 됨.

    Args:
        data: 파일 바이트 (업로드 등). filename과 쌍으로 사용.
        path: 디스크 상 파일 경로. PDF는 StrategyFactory 사용.
        filename: 확장자 판별용 (data 사용 시 필수). 예: "resume.pdf", "note.docx"

    Returns:
        추출된 전체 텍스트.

    Raises:
        ValueError: data/path 모두 없거나, 지원하지 않는 확장자.
        RuntimeError: 추출 실패 또는 필요한 라이브러리 없음.
    """
    ext = ""
    if path is not None:
        path = Path(path)
        if not path.exists():
            raise ValueError(f"파일이 없습니다: {path}")
        ext = (path.suffix or "").lower()
    elif data is not None and filename:
        ext = (Path(filename).suffix or "").lower()
    else:
        raise ValueError("data+filename 또는 path 중 하나를 지정해야 합니다.")

    if ext not in SUPPORTED_TEXT_EXTENSIONS:
        raise ValueError(
            f"지원하지 않는 확장자입니다: {ext}. 지원(텍스트): {', '.join(SUPPORTED_TEXT_EXTENSIONS)}"
        )

    if path is not None:
        if ext == ".pdf":
            return _extract_pdf_from_path(path)
        if ext == ".txt":
            raw = path.read_bytes()
            return _extract_txt_from_bytes(raw)
        if ext == ".docx":
            return _extract_docx_from_path(path)
        if ext == ".hwp":
            return _extract_hwp_from_path(path)

    # data + filename (이 시점에 data는 위 elif에서 설정됨)
    assert data is not None
    if ext == ".txt":
        return _extract_txt_from_bytes(data)
    if ext == ".pdf":
        return _extract_pdf_from_bytes(data)
    if ext == ".docx":
        return _extract_docx_from_bytes(data)
    if ext == ".hwp":
        return _extract_hwp_from_bytes(data)

    raise ValueError(f"처리할 수 없는 확장자: {ext}")


def _extract_xlsx_from_bytes(data: bytes) -> List[Dict[str, Any]]:
    """Excel(.xlsx): bytes → openpyxl로 첫 시트 읽기, 첫 행 헤더 → list[dict]."""
    try:
        from openpyxl import load_workbook  # type: ignore[import-untyped]
    except ImportError:
        raise RuntimeError("Excel(.xlsx) 추출을 위해 openpyxl이 필요합니다. pip install openpyxl")
    try:
        wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        ws = wb.active
        if ws is None:
            return []
        rows = list(ws.iter_rows(values_only=True))
        wb.close()
        if not rows:
            return []
        headers = [str(c or "").strip() or f"_col{i}" for i, c in enumerate(rows[0])]
        return [dict(zip(headers, (cell for cell in row))) for row in rows[1:]]
    except Exception as e:
        raise RuntimeError(f"Excel(.xlsx) 추출 실패: {e}")


def _extract_xlsx_from_path(path: Path) -> List[Dict[str, Any]]:
    """Excel(.xlsx): 경로 → openpyxl로 첫 시트 읽기."""
    try:
        from openpyxl import load_workbook  # type: ignore[import-untyped]
    except ImportError:
        raise RuntimeError("Excel(.xlsx) 추출을 위해 openpyxl이 필요합니다. pip install openpyxl")
    try:
        wb = load_workbook(path, read_only=True, data_only=True)
        ws = wb.active
        if ws is None:
            return []
        rows = list(ws.iter_rows(values_only=True))
        wb.close()
        if not rows:
            return []
        headers = [str(c or "").strip() or f"_col{i}" for i, c in enumerate(rows[0])]
        return [dict(zip(headers, (cell for cell in row))) for row in rows[1:]]
    except Exception as e:
        raise RuntimeError(f"Excel(.xlsx) 추출 실패: {e}")


def extract_excel_from_document(
    *,
    data: Optional[bytes] = None,
    path: Optional[Path] = None,
    filename: Optional[str] = None,
) -> List[Dict[str, Any]]:
    """문서에서 표 형식 추출 (Excel .xlsx).

    (data + filename) 또는 path 중 하나만 넘기면 됨.
    첫 시트, 첫 행을 헤더로 하여 list[dict] 반환.
    """
    ext = ""
    if path is not None:
        path = Path(path)
        if not path.exists():
            raise ValueError(f"파일이 없습니다: {path}")
        ext = (path.suffix or "").lower()
    elif data is not None and filename:
        ext = (Path(filename).suffix or "").lower()
    else:
        raise ValueError("data+filename 또는 path 중 하나를 지정해야 합니다.")

    if ext not in SUPPORTED_EXCEL_EXTENSIONS:
        raise ValueError(
            f"지원하지 않는 확장자입니다: {ext}. 지원(Excel): {', '.join(SUPPORTED_EXCEL_EXTENSIONS)}"
        )

    if path is not None:
        return _extract_xlsx_from_path(path)
    assert data is not None
    return _extract_xlsx_from_bytes(data)
